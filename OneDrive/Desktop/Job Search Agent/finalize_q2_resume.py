"""
Finalize Q2 Master Resume — incorporates:
  1. Approved Architect + B&C combo opening (kills the tricolon)
  2. Adds ChatGPT Paid Ads detail to Hyatt summary line
  3. Expands UPenn to full "double major" description
  4. Fixes awkward "surfaced through" phrasing in Masco line
  5. GTM QA bullet: "browser dev tools" → "browser console developer mode"
  6. MMM bullet: adds "(HTML)" for visualization layer specificity
Then re-exports DOCX + PDF.
"""
import shutil, os
from docx import Document

DOCX = r"C:\Users\alima\OneDrive\Desktop\Job Search Agent\Ali Mabsoute_Resume_Q2_Master.docx"
PDF  = r"C:\Users\alima\OneDrive\Desktop\Job Search Agent\Ali Mabsoute_Resume_Q2_Master.pdf"

# ── FINAL Q2 SUMMARY ─────────────────────────────────────────────────────────
# - Opens with "Architect" (approved)
# - B+C combo: role + scope in one em-dashed line
# - No tricolons
# - ChatGPT Paid Ads included
# - Full UPenn signal
# - "Manager" nowhere in first two sentences
Q2_SUMMARY_FINAL = (
    "Marketing Strategy & Analytics Architect with 15+ years building and governing "
    "marketing functions across Fortune 500 portfolios — from a $200M acquisition budget "
    "at Citigroup to a $178M revenue P&L at Hyatt. PMP + Six Sigma Black Belt. "
    "At Citigroup (AVP, U.S. Cards), directed two analysts across the U.S. Cards acquisition "
    "portfolio and rebuilt co-brand ROAS with a COVID-era Marketing Mix Model. "
    "At Masco, deployed an AI-driven personalization platform that cut CPA by 27% and drove "
    "$3.2M in incremental revenue through monthly C-suite readouts. "
    "At Hyatt, owns a $29M paid-media P&L; delivered a 40% ROAS lift on a 2% spend increase, "
    "led the proposal and CBA for the ChatGPT Paid Ads launch, and turned a multi-day "
    "C-suite reporting workflow into a single-prompt pipeline (95–98% faster). "
    "University of Pennsylvania — double major in Economics and Philosophy, Politics & "
    "Economics (PPE); M.S. Finance, Temple University."
)

def replace_para_text(para, new_text):
    """Overwrite paragraph runs while keeping first run's formatting."""
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run._r.getparent().remove(run._r)

def edit_run_text(para, old_substring, new_substring):
    """Find the run containing old_substring and patch in-place."""
    for run in para.runs:
        if old_substring in run.text:
            run.text = run.text.replace(old_substring, new_substring)
            return True
    # Fallback: rebuild the full paragraph text if substring spans runs
    full_text = para.text
    if old_substring in full_text:
        replace_para_text(para, full_text.replace(old_substring, new_substring))
        return True
    return False

doc = Document(DOCX)

# ── 1. Replace summary (para 3) ──────────────────────────────────────────────
summary_para = doc.paragraphs[3]
assert "Marketing Strategy & Analytics leader" in summary_para.text, \
    f"Summary mismatch: {summary_para.text[:80]}"
replace_para_text(summary_para, Q2_SUMMARY_FINAL)
print("[OK] Summary: Architect opening + ChatGPT detail + full UPenn + tricolon removed")

# ── 2. Fix Hyatt MMM bullet: add "(HTML)" ────────────────────────────────────
for p in doc.paragraphs:
    if "Full-Stack Marketing Mix Modeling" in p.text:
        if "visualization layer (HTML)" in p.text:
            print("[SKIP] MMM bullet already has (HTML)")
        elif edit_run_text(p, "visualization layer and Advanced Excel",
                              "visualization layer (HTML) and Advanced Excel"):
            print("[OK] MMM bullet: added (HTML) for visualization layer")
        else:
            print("[WARN] MMM bullet: couldn't patch visualization layer text")
        break

# ── 3. Fix GTM QA bullet wording ─────────────────────────────────────────────
for p in doc.paragraphs:
    if "Operational QA" in p.text:
        if edit_run_text(p, "browser dev tools", "browser console developer mode"):
            print("[OK] GTM QA bullet: browser dev tools -> browser console developer mode")
        else:
            print("[SKIP] GTM QA bullet: already correct or not found")
        break

# ── 4. Save ──────────────────────────────────────────────────────────────────
doc.save(DOCX)
print(f"[OK] DOCX saved: {DOCX}")

# ── 5. Export PDF via Word COM ───────────────────────────────────────────────
try:
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    abs_docx = os.path.abspath(DOCX)
    abs_pdf  = os.path.abspath(PDF)
    doc_com = word.Documents.Open(abs_docx)
    doc_com.SaveAs(abs_pdf, FileFormat=17)
    # Page count check
    page_count = doc_com.ComputeStatistics(2)  # wdStatisticPages
    doc_com.Close()
    word.Quit()
    print(f"[OK] PDF exported: {abs_pdf}")
    print(f"[OK] Page count: {page_count}")
except Exception as e:
    print(f"[WARN] PDF export failed: {e}")
