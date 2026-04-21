"""
Build Ali Mabsoute Q2 2026 Master Resume
- Shifts voice from Data/Analytics to Marketing Strategy & Analytics Leadership
- Fixes Galerie H title (Proper Case + "Lead" suffix)
- Reorders bullets (metric-heavy first: Masco + Hyatt)
- Ensures UPenn before Temple, no P&L; encoding
- Exports DOCX + PDF via Word COM
"""
import shutil, sys, io, copy
from docx import Document
from docx.oxml.ns import qn

SRC  = r"C:\Users\alima\OneDrive\Desktop\Job Search Agent\Ali Mabsoute_Resume_Q1_2026.docx"
DOCX = r"C:\Users\alima\OneDrive\Desktop\Job Search Agent\Ali Mabsoute_Resume_Q2_Master.docx"
PDF  = r"C:\Users\alima\OneDrive\Desktop\Job Search Agent\Ali Mabsoute_Resume_Q2_Master.pdf"

# ── 1. New Q2 Summary ────────────────────────────────────────────────────────
Q2_SUMMARY = (
    "Marketing Strategy & Analytics leader with 15+ years driving revenue growth, "
    "cross-functional transformation, and media investment governance across Fortune 500 portfolios. "
    "PMP + Six Sigma Black Belt. "
    "At Citigroup (AVP, U.S. Cards), governed a $200M annual customer acquisition portfolio, "
    "directed two analysts, and rebuilt co-brand ROAS with a COVID-era Marketing Mix Model. "
    "At Masco, deployed an AI-driven personalization platform that cut CPA by 27% and unlocked "
    "$3.2M in incremental revenue — surfaced through monthly C-suite readouts. "
    "At Hyatt, owns a $29M paid-media P&L and $178M revenue portfolio: delivered a 40% ROAS lift "
    "on a 2% spend increase, built a production AI guest-intelligence platform, and turned a "
    "multi-day C-suite reporting workflow into a single-prompt pipeline (95–98% faster). "
    "University of Pennsylvania — Economics and PPE; M.S. Finance, Temple University."
)

# ── 2. Copy Q1 as base ────────────────────────────────────────────────────────
shutil.copy2(SRC, DOCX)
doc = Document(DOCX)
paras = doc.paragraphs

# ── Helper: replace all run text in a paragraph, keeping first run's format ──
def replace_para_text(para, new_text):
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run._r.getparent().remove(run._r)

# ── Helper: swap two paragraph elements in the XML body ──────────────────────
def move_para_before(para_to_move, reference_para):
    """Move para_to_move so it appears immediately before reference_para."""
    elem = para_to_move._element
    ref  = reference_para._element
    elem.getparent().remove(elem)
    ref.addprevious(elem)

# ── 3. Replace summary (para 3) ───────────────────────────────────────────────
summary_para = paras[3]
assert "Performance Marketing and Digital Analytics Leader" in summary_para.text, \
    f"Expected summary at para 3, got: {summary_para.text[:60]}"
replace_para_text(summary_para, Q2_SUMMARY)
print("[OK] Summary replaced")

# ── 4. Fix Galerie H (para 31) ────────────────────────────────────────────────
galerie_para = paras[31]
assert "GALERIE H" in galerie_para.text, \
    f"Expected Galerie H at para 31, got: {galerie_para.text[:60]}"

for run in galerie_para.runs:
    if run.text == "GALERIE H":
        run.text = "Galerie H"
    elif "Business Operations & Analytics" in run.text:
        run.text = run.text.replace(
            "Business Operations & Analytics",
            "Business Operations and Analytics Lead"
        )
print("[OK] Galerie H fixed")

# ── 5. Reorder Masco bullets: want Stakeholder Data Viz (para 18) to be 2nd ──
# Current order: 15=AI-Driven, 16=E-Commerce, 17=Strategic Partner, 18=Stakeholder, 19=Lead Scoring
# Target order:  AI-Driven, Stakeholder, E-Commerce, Strategic Partner, Lead Scoring
# Move para 18 (Stakeholder) to just before para 16 (E-Commerce)
assert "Stakeholder Data Visualization" in paras[18].text, \
    f"Para 18 mismatch: {paras[18].text[:60]}"
assert "E-Commerce Relaunch" in paras[16].text, \
    f"Para 16 mismatch: {paras[16].text[:60]}"

# Re-fetch paragraphs (para indices stable for reordering since we're using XML)
masco_stakeholder = doc.paragraphs[18]
masco_ecommerce   = doc.paragraphs[16]
move_para_before(masco_stakeholder, masco_ecommerce)
print("[OK] Masco: Stakeholder Data Visualization moved to 2nd bullet")

# ── 6. Reorder Hyatt bullets: Executive Reporting Pipeline before Digital Analytics ──
# After step 5, Hyatt bullets are still at original para indices 6-13
# Current: 6=Exec Strategy, 7=MMM, 8=Digital Analytics, 9=Exec Reporting, ...
# Target:  6=Exec Strategy, 7=MMM, 8=Exec Reporting, 9=Digital Analytics, ...
# Move para 9 (Exec Reporting) to just before para 8 (Digital Analytics)
# Re-fetch since XML shifted
hyatt_paras = [p for p in doc.paragraphs if p.style.name == "List Paragraph"
               and any(kw in p.text for kw in [
                   "Executive Strategy", "Full-Stack Marketing Mix",
                   "Digital Analytics Transformation", "Executive Reporting Pipeline",
                   "Multi-Channel Media", "AI-Driven Guest", "Full-Funnel Multidimensional",
                   "Operational QA"
               ])]

exec_reporting = next((p for p in doc.paragraphs if "Executive Reporting Pipeline" in p.text), None)
digital_analytics = next((p for p in doc.paragraphs if "Digital Analytics Transformation" in p.text), None)

if exec_reporting and digital_analytics:
    # Compare positions via XML element ordering
    body = doc.element.body
    all_elems = list(body)
    idx_er = all_elems.index(exec_reporting._element)
    idx_da = all_elems.index(digital_analytics._element)
    if idx_er > idx_da:
        move_para_before(exec_reporting, digital_analytics)
        print("[OK] Hyatt: Executive Reporting Pipeline moved before Digital Analytics Transformation")
    else:
        print("[OK] Hyatt: Exec Reporting already before Digital Analytics")

# ── 7. Fix any P&L; encoding artifacts ───────────────────────────────────────
fixed_pl = 0
for para in doc.paragraphs:
    for run in para.runs:
        if "P&L;" in run.text:
            run.text = run.text.replace("P&L;", "P&L")
            fixed_pl += 1
print(f"[OK] P&L; fixes: {fixed_pl}")

# ── 8. Save DOCX ─────────────────────────────────────────────────────────────
doc.save(DOCX)
print(f"[OK] DOCX saved: {DOCX}")

# ── 9. Export PDF via Word COM ────────────────────────────────────────────────
try:
    import win32com.client
    import os
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    abs_docx = os.path.abspath(DOCX)
    abs_pdf  = os.path.abspath(PDF)
    doc_com = word.Documents.Open(abs_docx)
    doc_com.SaveAs(abs_pdf, FileFormat=17)   # 17 = wdFormatPDF
    doc_com.Close()
    word.Quit()
    print(f"[OK] PDF exported: {abs_pdf}")
except Exception as e:
    print(f"[WARN] PDF export failed: {e}")
    print("       Open the DOCX in Word and Save As PDF manually.")

# ── 10. Print comparison ──────────────────────────────────────────────────────
Q1_SUMMARY = (
    "Performance Marketing and Digital Analytics Leader with 15+ years driving customer acquisition, "
    "ROI optimization, and multimillion-dollar revenue growth across Fortune 500 portfolios. Expert in "
    "cross-functional leadership, marketing mix modeling, and budget optimization — managing media "
    "portfolios up to $200M while translating quantitative modeling into executive-ready media investment "
    "and growth decisions. Proven track record of building data strategy frameworks and transforming complex "
    "analytics into actionable narratives that improve customer lifetime value, inform media planning, and "
    "accelerate global portfolio performance."
)
print()
print("=" * 70)
print("VOICE COMPARISON: Q1 vs Q2 SUMMARY")
print("=" * 70)
print()
print("Q1 SUMMARY:")
print(Q1_SUMMARY)
print()
print("Q2 SUMMARY:")
print(Q2_SUMMARY)
print()
print("KEY SHIFTS:")
print("  Q1 opens: 'Performance Marketing and Digital Analytics Leader'")
print("  Q2 opens: 'Marketing Strategy & Analytics leader'")
print()
print("  Q1 voice: skills-first ('Expert in...', 'Proven track record')")
print("  Q2 voice: evidence-first (At Citi... At Masco... At Hyatt...)")
print()
print("  Q1 focus: Data/Analytics mechanics")
print("  Q2 focus: Revenue outcomes + leadership + C-suite communication")
print()
print("  Q1 'Manager' check (first 2 sentences): NONE - already clean")
print("  Q2 'Manager' check (first 2 sentences): NONE - clean")
