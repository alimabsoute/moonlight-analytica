"""
Trim Q2 to exactly 2 pages:
  1. Shorten summary - drop redundant "At Citigroup..." sentence (already in opener)
  2. Remove Hyatt Operational QA bullet (too tactical for Architect framing)
"""
import os
from docx import Document

DOCX = r"C:\Users\alima\OneDrive\Desktop\Job Search Agent\Ali Mabsoute_Resume_Q2_Master.docx"
PDF  = r"C:\Users\alima\OneDrive\Desktop\Job Search Agent\Ali Mabsoute_Resume_Q2_Master.pdf"

Q2_SUMMARY_TRIMMED = (
    "Marketing Strategy & Analytics Architect with 15+ years building and governing "
    "marketing functions across Fortune 500 portfolios — from a $200M acquisition budget "
    "at Citigroup to a $178M revenue P&L at Hyatt. PMP + Six Sigma Black Belt. "
    "At Masco, deployed an AI-driven personalization platform that cut CPA by 27% and drove "
    "$3.2M in incremental revenue through monthly C-suite readouts. "
    "At Hyatt, owns a $29M paid-media P&L; delivered a 40% ROAS lift on a 2% spend increase, "
    "led the proposal and CBA for the ChatGPT Paid Ads launch, and turned a multi-day "
    "C-suite reporting workflow into a single-prompt pipeline (95–98% faster). "
    "University of Pennsylvania — double major in Economics and Philosophy, Politics & "
    "Economics (PPE); M.S. Finance, Temple University."
)

def replace_para_text(para, new_text):
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run._r.getparent().remove(run._r)

doc = Document(DOCX)

# Trim summary
summary_para = doc.paragraphs[3]
assert "Marketing Strategy & Analytics Architect" in summary_para.text
replace_para_text(summary_para, Q2_SUMMARY_TRIMMED)
print("[OK] Summary trimmed")

# Remove Operational QA bullet
removed = False
for p in list(doc.paragraphs):
    if "Operational QA & Agency Management" in p.text:
        p._element.getparent().remove(p._element)
        removed = True
        print("[OK] Removed Operational QA bullet")
        break
if not removed:
    print("[WARN] Operational QA bullet not found")

doc.save(DOCX)
print(f"[OK] DOCX saved")

# Export PDF + page count
import win32com.client
word = win32com.client.Dispatch("Word.Application")
word.Visible = False
abs_docx = os.path.abspath(DOCX)
abs_pdf  = os.path.abspath(PDF)
d = word.Documents.Open(abs_docx)
d.SaveAs(abs_pdf, FileFormat=17)
pages = d.ComputeStatistics(2)
d.Close()
word.Quit()
print(f"[OK] PDF exported. Pages: {pages}")
